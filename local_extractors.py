from pathlib import Path


MAX_TEXT_CHARS = 18000


def extract_local_text(path, original_name):
    suffix = Path(original_name).suffix.lower()

    try:
        if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            text = _extract_xlsx(path)
        elif suffix == ".docx":
            text = _extract_docx(path)
        elif suffix == ".pdf":
            text = _extract_pdf(path)
        elif suffix in {".txt", ".csv", ".tsv"}:
            text = Path(path).read_text(encoding="utf-8", errors="ignore")
        else:
            text = ""
    except Exception as exc:
        return {
            "file_name": original_name,
            "source": "local",
            "status": "failed",
            "text": "",
            "error": str(exc),
        }

    return {
        "file_name": original_name,
        "source": "local",
        "status": "ok" if text.strip() else "empty",
        "text": _clip(text),
        "error": "",
    }


def _extract_xlsx(path):
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append("\t".join(values))
    workbook.close()
    return "\n".join(parts)


def _extract_docx(path):
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"[Table {table_index}]")
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_pdf(path):
    import pdfplumber

    parts = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            if text.strip():
                parts.append(f"[PDF page {page_index} text]")
                parts.append(text)
            for table_index, table in enumerate(tables, start=1):
                parts.append(f"[PDF page {page_index} table {table_index}]")
                for row in table:
                    parts.append("\t".join("" if cell is None else str(cell) for cell in row))
    return "\n".join(parts)


def _clip(text):
    text = text.strip()
    if len(text) <= MAX_TEXT_CHARS:
        return text
    return text[:MAX_TEXT_CHARS] + "\n...[local extraction clipped]"
